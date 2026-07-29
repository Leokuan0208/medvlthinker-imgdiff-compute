#!/usr/bin/env python3
"""Minimal markdown -> .docx converter for this repo's docs.
Handles: # / ## / ### headings, | tables |, - / * bullets, > blockquotes, ``` code blocks,
--- rules, and inline **bold** / *italic* / `code`. Usage: md2docx.py in.md out.docx
"""
import sys, re, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.environ.get("MD_BASE", ".")   # dir to resolve figs/ paths against
EMBED = os.environ.get("MD_EMBED_FIGS", "0") == "1"
_seen_figs = set()
def maybe_embed(doc, line):
    if not EMBED: return
    for m in re.findall(r'`?(figs/[\w/.-]+\.png)`?', line):
        p = os.path.join(BASE, m)
        if os.path.exists(p) and m not in _seen_figs:
            _seen_figs.add(m)
            try:
                doc.add_picture(p, width=Inches(5.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap.add_run(os.path.basename(m)); r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x6d,0x74,0x80)
            except Exception: pass

def add_inline(paragraph, text):
    # split on **bold**, `code`, *italic*  (process bold/code first)
    tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = paragraph.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        elif tok.startswith('*') and tok.endswith('*'):
            r = paragraph.add_run(tok[1:-1]); r.italic = True
        else:
            paragraph.add_run(tok)

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),hexcolor); tcPr.append(shd)
INDIGO=RGBColor(0x39,0x49,0xab); TEAL=RGBColor(0x00,0x69,0x5c)
def flush_table(doc, rows):
    rows = [r for r in rows if not re.match(r'^\s*\|[\s|:-]+\|\s*$', r)]  # drop separator row
    cells = [[c.strip() for c in r.strip().strip('|').split('|')] for r in rows]
    if not cells:
        return
    ncol = max(len(r) for r in cells)
    t = doc.add_table(rows=0, cols=ncol); t.style = 'Table Grid'
    for i, row in enumerate(cells):
        wcells = t.add_row().cells
        for j in range(ncol):
            wcells[j].paragraphs[0].clear()
            add_inline(wcells[j].paragraphs[0], row[j] if j < len(row) else '')
            wcells[j].paragraphs[0].runs and setattr(wcells[j].paragraphs[0].runs[0].font,'size',Pt(9.5))
            if i == 0:
                _shade(wcells[j], "3949AB")
                for run in wcells[j].paragraphs[0].runs:
                    run.bold = True; run.font.color.rgb = RGBColor(0xff,0xff,0xff)
            elif i % 2 == 1:
                _shade(wcells[j], "F2F5FB")

def convert(md_path, docx_path):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'; doc.styles['Normal'].font.size = Pt(10.5)
    lines = open(md_path, encoding='utf-8').read().split('\n')
    i = 0; tbl = []; incode = False; code = []
    def flush_code():
        if code:
            p = doc.add_paragraph(); r = p.add_run('\n'.join(code))
            r.font.name = 'Consolas'; r.font.size = Pt(8.5)
    while i < len(lines):
        ln = lines[i]
        if ln.strip().startswith('```'):
            if incode:
                flush_code(); code.clear(); incode = False
            else:
                incode = True
            i += 1; continue
        if incode:
            code.append(ln); i += 1; continue
        if ln.strip().startswith('|') and '|' in ln.strip()[1:]:
            tbl.append(ln); i += 1; continue
        elif tbl:
            flush_table(doc, tbl); tbl = []
        if re.match(r'^#{1,6}\s', ln):
            lvl = len(ln) - len(ln.lstrip('#')); txt = ln.lstrip('#').strip()
            h = doc.add_heading('', level=min(lvl, 4)); add_inline(h, txt)
            col = INDIGO if lvl <= 2 else TEAL
            for run in h.runs: run.font.color.rgb = col
            if lvl == 1:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in h.runs: run.font.size = Pt(20)
        elif ln.strip() in ('---', '***', '___'):
            doc.add_paragraph().add_run('_' * 60).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        elif ln.strip().startswith('>'):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
            add_inline(p, ln.strip().lstrip('>').strip());
            for r in p.runs: r.italic = True
        elif re.match(r'^\s*[-*]\s', ln):
            p = doc.add_paragraph(style='List Bullet'); add_inline(p, re.sub(r'^\s*[-*]\s', '', ln))
        elif re.match(r'^\s*\d+\.\s', ln):
            p = doc.add_paragraph(style='List Number'); add_inline(p, re.sub(r'^\s*\d+\.\s', '', ln))
        elif ln.strip() == '':
            pass
        else:
            p = doc.add_paragraph(); add_inline(p, ln); maybe_embed(doc, ln)
        i += 1
    if tbl: flush_table(doc, tbl)
    if code: flush_code()
    doc.save(docx_path)
    print('wrote', docx_path)

if __name__ == '__main__':
    convert(sys.argv[1], sys.argv[2])
